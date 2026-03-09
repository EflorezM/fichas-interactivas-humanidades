import streamlit as st
from docx import Document
import io

# Configuración de la página
st.set_page_config(page_title="Ficha de Religión - El judaísmo", page_icon="📜", layout="centered")

# Encabezado Oficial
st.markdown("### FICHA DE RELIGIÓN – SESIÓN 1° AÑO A - B - C")
st.markdown("**Tema:** El judaísmo: Cuna de las religiones monoteístas.")
st.markdown("**Prof:** Eduardo Florez Montero / Unidad 1")
st.markdown("---")

# Datos del estudiante
col1, col2, col3 = st.columns([2, 1, 1])
with col1:
    nombre = st.text_input("Estudiante:")
with col2:
    fecha = st.text_input("Fecha:")
with col3:
    st.info("Tiempo: 30min")

st.markdown("---")

# Secciones de contexto (Meta, Indicaciones, etc.)
with st.expander("🎯 META DE HOY", expanded=True):
    st.write("Reconocer aportes del judaísmo y practicar diálogo respetuoso sin prejuicios.")

with st.expander("📋 INDICACIONES", expanded=True):
    st.write("1) Resuelve TODO el Nivel 1 (Fácil).")
    st.write("2) Luego elige UN reto: Nivel 2 (Medio) o Nivel 3 (Difícil).")
    st.write("3) Si falta tiempo, completa en casa. No uses nombres reales ni datos personales.")

col_a, col_b = st.columns(2)
with col_a:
    st.success("**APOYOS (DUA)**\n- Línea de tiempo en pizarra digital.\n- Cuadro comparativo guiado.\n- Regla: escuchar sin interrumpir.")
with col_b:
    st.warning("**BANCO DE PALABRAS**\njudaísmo, monoteísmo, Abraham, alianza, Torá, tradición, respeto, diálogo")

st.markdown("### LECTURAS POR NIVELES (ELIGE UNA)")
tab1, tab2, tab3 = st.tabs(["Texto A (Fácil)", "Texto B (Medio)", "Texto C (Difícil)"])
with tab1:
    st.write("El judaísmo cree en un solo Dios y valora la alianza con Abraham.")
with tab2:
    st.write("Abraham es importante para judíos y cristianos. Su historia enseña fe y obediencia.")
with tab3:
    st.write("Conocer otras religiones ayuda a dialogar con respeto y evita prejuicios, sin perder mi identidad cristiana.")

st.markdown("---")

# NIVEL 1
st.subheader("NIVEL 1 (FÁCIL) – PARA TODOS")
st.caption("⚠️ Debes completar todas las preguntas de este nivel para poder generar tu archivo.")
q1 = st.text_input("1) Monoteísmo significa creer en:")
q2 = st.radio("2) Abraham es importante para:", ["judaísmo", "cristianismo", "ambos"], horizontal=True)

st.write("3) 2 valores de Abraham:")
col_v1, col_v2 = st.columns(2)
with col_v1:
    q3_1 = st.text_input("Valor 1:")
with col_v2:
    q3_2 = st.text_input("Valor 2:")

q4 = st.text_input("4) 1 regla para dialogar con respeto:")
st.info("🎨 *ZONA DE GRÁFICO / ORGANIZADOR:* Recuerda realizar tu dibujo en tu cuaderno o en una hoja aparte para mostrarlo en clase.")

st.markdown("---")

# NIVEL 2
st.subheader("NIVEL 2 (MEDIO) – RETO 1")
st.write("5) 2 semejanzas entre judaísmo y cristianismo:")
col_s1, col_s2 = st.columns(2)
with col_s1:
    q5_1 = st.text_input("Semejanza 1:")
with col_s2:
    q5_2 = st.text_input("Semejanza 2:")

q6 = st.text_input("6) 1 prejuicio que debo evitar:")

st.write("TABLA DE COMPARACIÓN (Describe brevemente una idea para cada punto)")
col_t1, col_t2, col_t3 = st.columns(3)
with col_t1:
    t_jud = st.text_area("Judaísmo (idea):")
with col_t2:
    t_cris = st.text_area("Cristianismo (idea):")
with col_t3:
    t_act = st.text_area("Actitud de respeto:")

st.markdown("---")

# NIVEL 3
st.subheader("NIVEL 3 (DIFÍCIL) – RETO 2")
st.markdown("🌟 **¡RETO FINAL!** Demuestra todo lo que has aprendido completando este nivel.")
q7 = st.text_area("7) En 5 líneas: ¿por qué el respeto es parte de valores cristianos?")
st.write("8) 3 acciones de convivencia con otras creencias:")
q8_1 = st.text_input("Acción 1:")
q8_2 = st.text_input("Acción 2:")
q8_3 = st.text_input("Acción 3:")

st.markdown("---")

# NUEVA SECCIÓN: VALORACIÓN DE LA FICHA
st.subheader("📊 Valoración de la Actividad")
st.write("Tu opinión nos ayuda a mejorar nuestras clases interactivas.")
val_funcional = st.slider("1. ¿Qué tan fácil y funcional te pareció usar esta ficha digital?", 1, 5, 5, help="1 = Muy difícil, 5 = Muy fácil")
val_interes = st.radio("2. ¿El tema y las actividades te parecieron interesantes para tu aprendizaje?", ["Sí, mucho", "Estuvo bien", "No mucho", "Nada interesante"], horizontal=True)

st.markdown("---")
st.caption("🔒 SEGURIDAD: Respeta tu privacidad. Usa ejemplos sin nombres reales ni datos personales.")

# Lógica para generar Word con validaciones
if st.button("Generar mi Evidencia en Word"):
    # Validación: Nombre y Nivel 1 obligatorio
    if not nombre.strip():
        st.error("⚠️ Por favor, ingresa tu nombre en la parte superior antes de descargar.")
    elif not q1.strip() or not q3_1.strip() or not q3_2.strip() or not q4.strip():
        st.error("⚠️ ¡Alto ahí! Para descargar tu evidencia, primero debes completar todas las preguntas del **NIVEL 1 (FÁCIL)**. ¡Tú puedes!")
    else:
        doc = Document()
        doc.add_heading('FICHA DE RELIGIÓN – SESIÓN 1° AÑO A - B - C', level=1)
        doc.add_paragraph('Tema: El judaísmo: Cuna de las religiones monoteístas.')
        doc.add_paragraph('Prof: Eduardo Florez Montero / Unidad 1')
        doc.add_paragraph(f'Estudiante: {nombre} | Fecha: {fecha}')
        doc.add_paragraph('---')
        
        doc.add_heading('NIVEL 1 (FÁCIL)', level=2)
        doc.add_paragraph(f'1) Monoteísmo significa creer en: {q1}')
        doc.add_paragraph(f'2) Abraham es importante para: {q2}')
        doc.add_paragraph(f'3) 2 valores de Abraham: 1) {q3_1} | 2) {q3_2}')
        doc.add_paragraph(f'4) 1 regla para dialogar con respeto: {q4}')
        
        doc.add_heading('NIVEL 2 (MEDIO) - RETO 1', level=2)
        doc.add_paragraph(f'5) 2 semejanzas: 1) {q5_1} | 2) {q5_2}')
        doc.add_paragraph(f'6) 1 prejuicio que debo evitar: {q6}')
        doc.add_paragraph(f'TABLA: Judaísmo: {t_jud} | Cristianismo: {t_cris} | Actitud: {t_act}')
        
        doc.add_heading('NIVEL 3 (DIFÍCIL) - RETO 2', level=2)
        doc.add_paragraph(f'7) Por qué el respeto es parte de valores cristianos: {q7}')
        doc.add_paragraph(f'8) 3 acciones de convivencia: 1) {q8_1} | 2) {q8_2} | 3) {q8_3}')
        
        # AGREGAR LA VALORACIÓN AL DOCUMENTO
        doc.add_heading('Valoración de la Actividad', level=2)
        doc.add_paragraph(f'Funcionalidad de la ficha (1 a 5 estrellas): {val_funcional}')
        doc.add_paragraph(f'Interés en el aprendizaje: {val_interes}')
        
        bio = io.BytesIO()
        doc.save(bio)
        
        # Mensaje condicional dependiendo de si hicieron el nivel 3
        if q7.strip() and q8_1.strip():
            st.balloons()
            st.success("🌟 ¡Increíble! Has completado hasta el Nivel 3. ¡Excelente trabajo!")
        else:
            st.success("¡Tu archivo está listo para entregar! 🚀 Te reto a que, si aún tienes tiempo, regreses e intentes completar el Nivel 2 y Nivel 3 para un mayor aprendizaje.")
            
        st.download_button(
            label="📥 Descargar resolución en .docx",
            data=bio.getvalue(),
            file_name=f"Ficha_Religion_{nombre.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
