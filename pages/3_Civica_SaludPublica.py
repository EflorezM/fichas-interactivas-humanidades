import streamlit as st
from docx import Document
import io

# Configuración de la página
st.set_page_config(page_title="Ficha Cívica - Salud Pública", page_icon="⚖️", layout="centered")

# Encabezado Oficial [cite: 111, 113, 114]
st.markdown("### FICHA DE CÍVICA – SESIÓN 2° AÑO A - B - C")
st.markdown("**Tema:** Derecho a la Vida y la Salud Pública")
st.markdown("**Prof:** Eduardo Florez Montero / Unidad 1")
st.markdown("---")

# Datos del estudiante [cite: 115]
col1, col2, col3 = st.columns([2, 1, 1])
with col1:
    nombre = st.text_input("Estudiante:")
with col2:
    fecha = st.text_input("Fecha:")
with col3:
    st.info("Tiempo: 30min clase")

st.markdown("---")

# Secciones de contexto [cite: 116, 117, 118, 119, 120, 121]
with st.expander("🎯 META DE HOY", expanded=True):
    st.write("Relacionar el derecho a la vida con la salud pública y proponer normas de prevención e información responsable.")

with st.expander("📋 INDICACIONES", expanded=True):
    st.write("1) Resuelve TODO el Nivel 1 (Fácil).")
    st.write("2) Luego elige UN reto: Nivel 2 (Medio) o Nivel 3 (Difícil).")
    st.write("3) Si falta tiempo, completa en casa. No uses nombres reales.")

# Apoyos y Banco de Palabras [cite: 122, 123, 124, 125, 126, 127]
col_a, col_b = st.columns(2)
with col_a:
    st.success("**APOYOS (DUA)**\n- Ejemplos de prevención en pizarra digital.\n- Plantilla “medida–beneficio–responsabilidad”.\n- Puedes responder con viñetas.")
with col_b:
    st.warning("**BANCO DE PALABRAS**\nsalud pública, prevención, higiene, vacuna, información, verificar, responsabilidad, comunidad")

# Lecturas [cite: 128, 129, 130, 131, 132, 133, 134]
st.markdown("### LECTURAS POR NIVELES (ELIGE UNA)")
tab1, tab2, tab3 = st.tabs(["Texto A (Fácil)", "Texto B (Medio)", "Texto C (Difícil)"])
with tab1:
    st.write("La salud pública cuida la vida de todas las personas. La prevención ayuda a evitar enfermedades.")
with tab2:
    st.write("Compartir información responsable en redes también protege la vida.")
with tab3:
    st.write("La desinformación en salud puede poner vidas en riesgo. Verificar fuentes es clave.")

st.markdown("---")

# NIVEL 1 [cite: 135, 136, 137, 138, 139]
st.subheader("NIVEL 1 (FÁCIL) – PARA TODOS")
st.caption("⚠️ Debes completar todas las preguntas de este nivel para poder generar tu archivo.")

st.write("1) 2 acciones de prevención:")
col_n1a, col_n1b = st.columns(2)
with col_n1a:
    q1_1 = st.text_input("Acción 1:")
with col_n1b:
    q1_2 = st.text_input("Acción 2:")

q2 = text_input_2 = st.text_input("2) Difundir información falsa de salud puede...")
q3 = st.radio("3) Antes de compartir una noticia debo:", ["verificar", "reenviar sin leer", "insultar"], horizontal=True)
q4 = st.text_input("4) 1 compromiso de salud en el aula:")

st.info("🎨 *ORGANIZADOR GRÁFICO:* Recuerda realizar tu dibujo en tu cuaderno o en una hoja aparte para mostrarlo en clase. [cite: 140]")

st.markdown("---")

# NIVEL 2 [cite: 142, 143, 144, 145, 146]
st.subheader("NIVEL 2 (MEDIO) – RETO 1")
st.write("5) Caso: “Las vacunas hacen daño” (sin pruebas). 2 preguntas para verificar:")
q5_1 = st.text_input("Pregunta 1 para verificar:")
q5_2 = st.text_input("Pregunta 2 para verificar:")

st.write("6) 2 normas para compartir información de salud:")
q6_1 = st.text_input("Norma 1:")
q6_2 = st.text_input("Norma 2:")

st.write("TABLA DE COMPARACIÓN (Describe 3 medidas) [cite: 147, 148]")
# Fila 1
col_t1a, col_t1b, col_t1c, col_t1d = st.columns(4)
with col_t1a:
    t1_med = st.text_input("Medida 1:")
with col_t1b:
    t1_ben = st.text_input("Beneficio 1:")
with col_t1c:
    t1_res = st.text_input("Responsabilidad 1:")
with col_t1d:
    t1_eje = st.text_input("Ejemplo en redes 1:")
# Fila 2
col_t2a, col_t2b, col_t2c, col_t2d = st.columns(4)
with col_t2a:
    t2_med = st.text_input("Medida 2:")
with col_t2b:
    t2_ben = st.text_input("Beneficio 2:")
with col_t2c:
    t2_res = st.text_input("Responsabilidad 2:")
with col_t2d:
    t2_eje = st.text_input("Ejemplo en redes 2:")
# Fila 3
col_t3a, col_t3b, col_t3c, col_t3d = st.columns(4)
with col_t3a:
    t3_med = st.text_input("Medida 3:")
with col_t3b:
    t3_ben = st.text_input("Beneficio 3:")
with col_t3c:
    t3_res = st.text_input("Responsabilidad 3:")
with col_t3d:
    t3_eje = st.text_input("Ejemplo en redes 3:")

st.markdown("---")

# NIVEL 3 [cite: 149, 150, 151, 152]
st.subheader("NIVEL 3 (DIFÍCIL) – RETO 2")
st.markdown("🌟 **¡RETO FINAL!** Demuestra todo lo que has aprendido completando este nivel.")
q7 = st.text_area("7) Mensaje de campaña (2 líneas):")
q8 = st.text_area("8) En 5 líneas: ¿por qué la información responsable protege la vida?")

st.markdown("---")

# VALORACIÓN DE LA FICHA
st.subheader("📊 Valoración de la Actividad")
st.write("Tu opinión nos ayuda a mejorar nuestras clases interactivas.")
val_funcional = st.slider("1. ¿Qué tan fácil y funcional te pareció usar esta ficha digital?", 1, 5, 5)
val_interes = st.radio("2. ¿El tema y las actividades te parecieron interesantes para tu aprendizaje?", ["Sí, mucho", "Estuvo bien", "No mucho", "Nada interesante"], horizontal=True)

st.markdown("---")
st.caption("🔒 SEGURIDAD : No compartas datos personales. Usa ejemplos sin nombres reales. [cite: 156]")

# Lógica para generar Word
if st.button("Generar mi Evidencia en Word"):
    if not nombre.strip():
        st.error("⚠️ Por favor, ingresa tu nombre en la parte superior antes de descargar.")
    elif not q1_1.strip() or not q1_2.strip() or not q2.strip() or not q4.strip():
        st.error("⚠️ ¡Alto ahí! Para descargar tu evidencia, primero debes completar todas las preguntas del **NIVEL 1 (FÁCIL)**.")
    else:
        doc = Document()
        doc.add_heading('FICHA DE CÍVICA – SESIÓN 2° AÑO A - B - C', level=1)
        doc.add_paragraph('Tema: Derecho a la Vida y la Salud Pública')
        doc.add_paragraph('Prof: Eduardo Florez Montero / Unidad 1')
        doc.add_paragraph(f'Estudiante: {nombre} | Fecha: {fecha}')
        doc.add_paragraph('---')
        
        doc.add_heading('NIVEL 1 (FÁCIL)', level=2)
        doc.add_paragraph(f'1) Acciones de prevención: 1) {q1_1} | 2) {q1_2}')
        doc.add_paragraph(f'2) Difundir info falsa puede: {q2}')
        doc.add_paragraph(f'3) Antes de compartir debo: {q3}')
        doc.add_paragraph(f'4) Compromiso de salud: {q4}')
        
        doc.add_heading('NIVEL 2 (MEDIO) - RETO 1', level=2)
        doc.add_paragraph(f'5) Caso vacunas: 1) {q5_1} | 2) {q5_2}')
        doc.add_paragraph(f'6) Normas info salud: 1) {q6_1} | 2) {q6_2}')
        doc.add_paragraph(f'TABLA 1: Med: {t1_med} | Ben: {t1_ben} | Res: {t1_res} | Eje: {t1_eje}')
        doc.add_paragraph(f'TABLA 2: Med: {t2_med} | Ben: {t2_ben} | Res: {t2_res} | Eje: {t2_eje}')
        doc.add_paragraph(f'TABLA 3: Med: {t3_med} | Ben: {t3_ben} | Res: {t3_res} | Eje: {t3_eje}')
        
        doc.add_heading('NIVEL 3 (DIFÍCIL) - RETO 2', level=2)
        doc.add_paragraph(f'7) Mensaje de campaña: {q7}')
        doc.add_paragraph(f'8) Por qué la info responsable protege la vida: {q8}')
        
        doc.add_heading('Valoración de la Actividad', level=2)
        doc.add_paragraph(f'Funcionalidad de la ficha (1 a 5 estrellas): {val_funcional}')
        doc.add_paragraph(f'Interés en el aprendizaje: {val_interes}')
        
        bio = io.BytesIO()
        doc.save(bio)
        
        if q7.strip() and q8.strip():
            st.balloons()
            st.success("🌟 ¡Increíble! Has completado hasta el Nivel 3. ¡Excelente trabajo!")
        else:
            st.success("¡Tu archivo está listo para entregar! 🚀 Te reto a que regreses e intentes completar el Nivel 2 y Nivel 3.")
            
        st.download_button(
            label="📥 Descargar resolución en .docx",
            data=bio.getvalue(),
            file_name=f"Ficha_Civica_{nombre.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
