import streamlit as st
from docx import Document
import io

# Configuración de la página
st.set_page_config(page_title="Ficha PFRH - Emociones", page_icon="🧠", layout="centered")

# Encabezado Oficial
st.markdown("### FICHA DE PFRH – SESIÓN 3° AÑO A - B - C")
st.markdown("**Tema:** Las emociones que experimentamos.")
st.markdown("**Prof:** Eduardo Florez Montero / Unidad 1")
st.markdown("---")

# Datos del estudiante
col1, col2, col3 = st.columns([2, 1, 1])
with col1:
    nombre = st.text_input("Estudiante:")
with col2:
    fecha = st.text_input("Fecha:")
with col3:
    st.info("Tiempo: 25min clase")

st.markdown("---")

# Secciones de contexto (Meta, Indicaciones, etc.)
with st.expander("🎯 META DE HOY", expanded=True):
    st.write("Reconocer emoción–pensamiento–situación y expresar emociones con respeto.")

with st.expander("📋 INDICACIONES", expanded=True):
    st.write("1) Resuelve TODO el Nivel 1 (Fácil).")
    st.write("2) Luego elige UN reto: Nivel 2 (Medio) o Nivel 3 (Difícil).")
    st.write("3) Si falta tiempo, completa en casa. No uses nombres reales ni datos personales.")

col_a, col_b = st.columns(2)
with col_a:
    st.success("**APOYOS (DUA)**\n- Rueda de emociones en pizarra digital.\n- Plantilla: emoción → pensamiento → situación → respuesta.\n- Puedes usar palabras clave.")
with col_b:
    st.warning("**BANCO DE PALABRAS**\nemociones, pensamiento, situación, enojo, tristeza, miedo, alegría, vergüenza, calma")

st.markdown("### LECTURAS POR NIVELES (ELIGE UNA)")
tab1, tab2, tab3 = st.tabs(["Texto A (Fácil)", "Texto B (Medio)", "Texto C (Difícil)"])
with tab1:
    st.write("Las emociones son señales. No son malas: me ayudan a entender lo que necesito.")
with tab2:
    st.write("A veces pienso algo y eso cambia mi emoción. Si cambio el pensamiento, puedo responder mejor.")
with tab3:
    st.write("En redes, un comentario puede activar emociones fuertes. Respirar y responder con respeto evita conflictos.")

st.markdown("---")

# NIVEL 1
st.subheader("NIVEL 1 (FÁCIL) – PARA TODOS")
st.caption("⚠️ Debes completar todas las preguntas de este nivel para poder generar tu archivo.")

q1 = st.selectbox("1) Marca tu emoción de hoy:", ["", "alegría", "enojo", "tristeza", "miedo", "vergüenza", "calma"])
q2 = st.text_input("2) Completa: Cuando siento enojo, mi cuerpo...")
q3 = st.text_area("3) Escribe 1 pensamiento que suele aparecer (sin nombres):")
q4 = st.text_area("4) Escribe 1 forma respetuosa de expresar esa emoción:")

st.info("🎨 *ZONA DE GRÁFICO / ORGANIZADOR:* Recuerda realizar tu dibujo en tu cuaderno o en una hoja aparte para mostrarlo en clase.")

st.markdown("---")

# NIVEL 2
st.subheader("NIVEL 2 (MEDIO) – RETO 1")
st.write("5) Caso: Me dejan en visto en el chat.")
col_c1, col_c2, col_c3 = st.columns(3)
with col_c1:
    q5_emo = st.text_input("Emoción:")
with col_c2:
    q5_pen = st.text_input("Pensamiento:")
with col_c3:
    q5_res = st.text_input("Respuesta respetuosa:")

st.write("TABLA DE COMPARACIÓN (Describe 3 situaciones)")
# Fila 1
col_t1a, col_t1b, col_t1c, col_t1d = st.columns(4)
with col_t1a:
    t1_sit = st.text_input("Situación 1:")
with col_t1b:
    t1_emo = st.text_input("Emoción 1:")
with col_t1c:
    t1_pen = st.text_input("Pensamiento 1:")
with col_t1d:
    t1_res = st.text_input("Respuesta 1:")
# Fila 2
col_t2a, col_t2b, col_t2c, col_t2d = st.columns(4)
with col_t2a:
    t2_sit = st.text_input("Situación 2:")
with col_t2b:
    t2_emo = st.text_input("Emoción 2:")
with col_t2c:
    t2_pen = st.text_input("Pensamiento 2:")
with col_t2d:
    t2_res = st.text_input("Respuesta 2:")
# Fila 3
col_t3a, col_t3b, col_t3c, col_t3d = st.columns(4)
with col_t3a:
    t3_sit = st.text_input("Situación 3:")
with col_t3b:
    t3_emo = st.text_input("Emoción 3:")
with col_t3c:
    t3_pen = st.text_input("Pensamiento 3:")
with col_t3d:
    t3_res = st.text_input("Respuesta 3:")

st.markdown("---")

# NIVEL 3
st.subheader("NIVEL 3 (DIFÍCIL) – RETO 2")
st.markdown("🌟 **¡RETO FINAL!** Demuestra todo lo que has aprendido completando este nivel.")
q6 = st.text_area("6) Escribe 5 líneas: ¿qué pasa si guardo emociones y exploto en redes?")
st.write("7) Diseña 2 frases de autocuidado:")
q7_1 = st.text_input("Frase 1:")
q7_2 = st.text_input("Frase 2:")

st.markdown("---")

# VALORACIÓN DE LA FICHA
st.subheader("📊 Valoración de la Actividad")
st.write("Tu opinión nos ayuda a mejorar nuestras clases interactivas.")
val_funcional = st.slider("1. ¿Qué tan fácil y funcional te pareció usar esta ficha digital?", 1, 5, 5, help="1 = Muy difícil, 5 = Muy fácil")
val_interes = st.radio("2. ¿El tema y las actividades te parecieron interesantes para tu aprendizaje?", ["Sí, mucho", "Estuvo bien", "No mucho", "Nada interesante"], horizontal=True)

st.markdown("---")
st.caption("🔒 SEGURIDAD: Respeta tu privacidad. Usa ejemplos sin nombres reales ni datos personales.")

# Lógica para generar Word
if st.button("Generar mi Evidencia en Word"):
    if not nombre.strip():
        st.error("⚠️ Por favor, ingresa tu nombre en la parte superior antes de descargar.")
    elif q1 == "" or not q2.strip() or not q3.strip() or not q4.strip():
        st.error("⚠️ ¡Alto ahí! Para descargar tu evidencia, primero debes completar todas las preguntas del **NIVEL 1 (FÁCIL)**. ¡Tú puedes!")
    else:
        doc = Document()
        doc.add_heading('FICHA DE PFRH – SESIÓN 3° AÑO A - B - C', level=1)
        doc.add_paragraph('Tema: Las emociones que experimentamos.')
        doc.add_paragraph('Prof: Eduardo Florez Montero / Unidad 1')
        doc.add_paragraph(f'Estudiante: {nombre} | Fecha: {fecha}')
        doc.add_paragraph('---')
        
        doc.add_heading('NIVEL 1 (FÁCIL)', level=2)
        doc.add_paragraph(f'1) Emoción de hoy: {q1}')
        doc.add_paragraph(f'2) Cuando siento enojo, mi cuerpo: {q2}')
        doc.add_paragraph(f'3) Pensamiento que suele aparecer: {q3}')
        doc.add_paragraph(f'4) Forma respetuosa de expresar: {q4}')
        
        doc.add_heading('NIVEL 2 (MEDIO) - RETO 1', level=2)
        doc.add_paragraph(f'5) Caso (En visto): Emoción: {q5_emo} | Pensamiento: {q5_pen} | Respuesta: {q5_res}')
        doc.add_paragraph(f'TABLA 1: Sit: {t1_sit} | Emo: {t1_emo} | Pen: {t1_pen} | Res: {t1_res}')
        doc.add_paragraph(f'TABLA 2: Sit: {t2_sit} | Emo: {t2_emo} | Pen: {t2_pen} | Res: {t2_res}')
        doc.add_paragraph(f'TABLA 3: Sit: {t3_sit} | Emo: {t3_emo} | Pen: {t3_pen} | Res: {t3_res}')
        
        doc.add_heading('NIVEL 3 (DIFÍCIL) - RETO 2', level=2)
        doc.add_paragraph(f'6) Qué pasa si guardo emociones: {q6}')
        doc.add_paragraph(f'7) Frases autocuidado: 1) {q7_1} | 2) {q7_2}')
        
        doc.add_heading('Valoración de la Actividad', level=2)
        doc.add_paragraph(f'Funcionalidad de la ficha (1 a 5 estrellas): {val_funcional}')
        doc.add_paragraph(f'Interés en el aprendizaje: {val_interes}')
        
        bio = io.BytesIO()
        doc.save(bio)
        
        if q6.strip() and q7_1.strip():
            st.balloons()
            st.success("🌟 ¡Increíble! Has completado hasta el Nivel 3. ¡Excelente trabajo!")
        else:
            st.success("¡Tu archivo está listo para entregar! 🚀 Te reto a que, si aún tienes tiempo, regreses e intentes completar el Nivel 2 y Nivel 3.")
            
        st.download_button(
            label="📥 Descargar resolución en .docx",
            data=bio.getvalue(),
            file_name=f"Ficha_PFRH_{nombre.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
